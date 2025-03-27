from bot.models import Presentations, User  # –ò–º–ø–æ—Ä—Ç –≤–∞—à–µ–π –º–æ–¥–µ–ª–∏ –≤ Django
from docx import Document
from docx.shared import Inches
import os
from datetime import datetime
from django.conf import settings

from bot import bot
from bot.handlers.admin import admin_notification
from bot.keyboard import START_BUTTONS, WATER_SUPPLY, AGGREMENT_BUTTON, FLOOR_BUTTONS, TYPE_AREA, TYPE_RENT, \
    SKIP_BUTTON


def start_message(message):
    bot.clear_step_handler_by_chat_id(message.chat.id)
    user = User.objects.update_or_create(
        user_id=message.from_user.id,
        defaults={"user_id": message.from_user.id, "username": message.from_user.username}
    )
    bot.send_message(message.chat.id, "–ó–¥—Ä–∞–≤—Å—Ç–≤—É–π—Ç–µ! –Ø –ø–æ–º–æ–≥—É —Å–æ–±—Ä–∞—Ç—å –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ –ø–æ–º–µ—â–µ–Ω–∏–∏",
                     reply_markup=START_BUTTONS,
                     )


def exit_handler(call):
    pres = Presentations.objects.filter(user=User.objects.get(user_id=str(call.message.chat.id))).last()
    pres.delete()
    bot.clear_step_handler_by_chat_id(chat_id=call.message.chat.id)
    start_message(message=call.message)


def ask_question(call):
    pres = Presentations.objects.create(user=User.objects.get(user_id=call.message.chat.id))
    pres.save()

    def check_exit(message) -> bool:
        if message.text == "–ü—Ä–µ–∫—Ä–∞—Ç–∏—Ç—å —Å–æ–∑–¥–∞–Ω–∏–µ –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–∏—è":
            pres.delete()
            bot.clear_step_handler_by_chat_id(chat_id=message.chat.id)
            return True
        return False

    def create_pres(message):
        if check_exit(message):
            return
        # create_presentation(object=pres)
        bot.send_message(
            text="‚úâ –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –ø–µ—Ä–µ–¥–∞–Ω–∞ –∞–≥–µ–Ω—Ç—É\n\nüìû –ù–∞—à–∏ –∫–æ–Ω—Ç–∞–∫—Ç—ã –¥–ª—è —Å–≤—è–∑–∏: +7 933 481 00 01",
            chat_id=message.chat.id,
        )
        admin_notification(pres, message.chat.id)
        return

    def agreement(id_):
        msg = bot.send_message(
            text="–î–ª—è –ø—Ä–æ–¥–æ–ª–∂–µ–Ω–∏—è –Ω—É–∂–Ω–æ –ø–æ–¥—Ç–≤–µ—Ä–¥–∏—Ç—å —Å–≤–æ–µ —Å–æ–≥–ª–∞—Å–∏–µ –Ω–∞ –æ–±—Ä–∞–±–æ—Ç–∫—É –¥–∞–Ω–Ω—ã—Ö, –Ω–∞–∂–∞–≤ –Ω–∞ –∫–Ω–æ–ø–∫—É –Ω–∏–∂–µ",
            chat_id=id_,
            reply_markup=AGGREMENT_BUTTON,
        )
        bot.register_next_step_handler(msg, create_pres)

    def register_number(message):
        if check_exit(message):
            return
        pres.contact += message.text
        pres.save()
        return agreement(message.chat.id)

    def get_number(id_):
        msg = bot.send_message(
            text="–¢–µ–ª–µ—Ñ–æ–Ω –¥–ª—è —Å–≤—è–∑–∏", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_number)

    def register_contact(message):
        if check_exit(message):
            return
        pres.contact = f"{message.text}\n"
        pres.save()
        return get_number(message.chat.id)

    def get_contact(id_):
        msg = bot.send_message(
            text="–ö–∞–∫ –∫ –≤–∞–º –æ–±—Ä–∞—â–∞—Ç—å—Å—è (–§–ò–û)", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_contact)

    def register_additives(message):
        if check_exit(message):
            return
        if message.text == "–ü–æ–º–µ—Ç–∫–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤—É—é—Ç":
            return get_contact(message.chat.id)

        try:
            pres.additives = message.text
        except Exception as e:
            print(e)
        pres.save()
        return get_contact(message.chat.id)

    def get_additives(id_):
        msg = bot.send_message(
            text="–ï—Å–ª–∏ —É –≤–∞—Å –æ—Å—Ç–∞–ª–∏—Å—å –∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–∏ –∏ –ø–æ–∂–µ–ª–∞–Ω–∏—è - –Ω–∞–ø–∏—à–∏—Ç–µ –∏—Ö —Å—é–¥–∞\n\n–ï—Å–ª–∏ –ø–æ–º–µ—Ç–æ–∫ –Ω–µ—Ç - –Ω–∞–∂–º–∏—Ç–µ –Ω–∞ –∫–Ω–æ–ø–∫—É",
            chat_id=id_, reply_markup=SKIP_BUTTON
        )
        bot.register_next_step_handler(msg, register_additives)

    def register_photo_inside(message):
        if check_exit(message):
            return
        try:
            if message.photo:
                # –ü–æ–ª—É—á–∞–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ —Ñ–æ—Ç–æ (–±–µ—Ä–µ–º –ø–æ—Å–ª–µ–¥–Ω–µ–µ, —Ç.–∫. –æ–Ω–æ —Å–∞–º–æ–≥–æ –≤—ã—Å–æ–∫–æ–≥–æ –∫–∞—á–µ—Å—Ç–≤–∞)
                file_info = bot.get_file(message.photo[-1].file_id)
                downloaded_file = bot.download_file(file_info.file_path)

                # –§–æ—Ä–º–∏—Ä—É–µ–º –ø—É—Ç—å –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è
                file_name = f"inside_{message.chat.id}_{message.message_id}.jpg"
                file_path = os.path.join(settings.MEDIA_ROOT, 'temp', file_name)

                # –°–æ–∑–¥–∞–µ–º –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏–∏ –µ—Å–ª–∏ –∏—Ö –Ω–µ—Ç
                os.makedirs(os.path.dirname(file_path), exist_ok=True)

                # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ñ–∞–π–ª
                with open(file_path, 'wb') as new_file:
                    new_file.write(downloaded_file)

                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –ø—É—Ç—å –≤ –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö
                pres.photo_inside = file_path
                pres.save()
            else:
                bot.reply_to(message, "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –æ—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏—é")
                return get_photo_inside(message.chat.id)
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ —Ñ–æ—Ç–æ –≤–Ω—É—Ç—Ä–∏: {e}")
            bot.reply_to(message, "–ü—Ä–æ–∏–∑–æ—à–ª–∞ –æ—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ —Ñ–æ—Ç–æ. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ –µ—â–µ —Ä–∞–∑.")
            return get_photo_inside(message.chat.id)
        return get_additives(message.chat.id)

    def get_photo_inside(id_):
        msg = bot.send_message(
            text="–û—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ –≤–Ω—É—Ç—Ä–∏ –ø–æ–º–µ—â–µ–Ω–∏—è (–æ—Ç–ø—Ä–∞–≤—å—Ç–µ –∫–∞–∫ —Ñ–æ—Ç–æ, –Ω–µ –∫–∞–∫ —Ñ–∞–π–ª)", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_photo_inside)

    def register_photo_outside(message):
        if check_exit(message):
            return
        try:
            if message.photo:
                # –ü–æ–ª—É—á–∞–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ —Ñ–æ—Ç–æ (–±–µ—Ä–µ–º –ø–æ—Å–ª–µ–¥–Ω–µ–µ, —Ç.–∫. –æ–Ω–æ —Å–∞–º–æ–≥–æ –≤—ã—Å–æ–∫–æ–≥–æ –∫–∞—á–µ—Å—Ç–≤–∞)
                file_info = bot.get_file(message.photo[-1].file_id)
                downloaded_file = bot.download_file(file_info.file_path)

                # –§–æ—Ä–º–∏—Ä—É–µ–º –ø—É—Ç—å –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è
                file_name = f"outside_{message.chat.id}_{message.message_id}.jpg"
                file_path = os.path.join(settings.MEDIA_ROOT, 'temp', file_name)

                # –°–æ–∑–¥–∞–µ–º –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏–∏ –µ—Å–ª–∏ –∏—Ö –Ω–µ—Ç
                os.makedirs(os.path.dirname(file_path), exist_ok=True)

                # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ñ–∞–π–ª
                with open(file_path, 'wb') as new_file:
                    new_file.write(downloaded_file)

                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –ø—É—Ç—å –≤ –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö
                pres.photo_outside = file_path
                pres.save()
            else:
                bot.reply_to(message, "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –æ—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏—é")
                return get_photo_outside(message.chat.id)
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ —Ñ–æ—Ç–æ —Å–Ω–∞—Ä—É–∂–∏: {e}")
            bot.reply_to(message, "–ü—Ä–æ–∏–∑–æ—à–ª–∞ –æ—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ —Ñ–æ—Ç–æ. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ –µ—â–µ —Ä–∞–∑.")
            return get_photo_outside(message.chat.id)
        return get_photo_inside(message.chat.id)

    def get_photo_outside(id_):
        msg = bot.send_message(
            text="–û—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ —Å–Ω–∞—Ä—É–∂–∏ –ø–æ–º–µ—â–µ–Ω–∏—è (–æ—Ç–ø—Ä–∞–≤—å—Ç–µ –∫–∞–∫ —Ñ–æ—Ç–æ, –Ω–µ –∫–∞–∫ —Ñ–∞–π–ª)", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_photo_outside)

    def register_plan(message):
        if check_exit(message):
            return
        try:
            if message.photo:
                # –ü–æ–ª—É—á–∞–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ —Ñ–æ—Ç–æ (–±–µ—Ä–µ–º –ø–æ—Å–ª–µ–¥–Ω–µ–µ, —Ç.–∫. –æ–Ω–æ —Å–∞–º–æ–≥–æ –≤—ã—Å–æ–∫–æ–≥–æ –∫–∞—á–µ—Å—Ç–≤–∞)
                file_info = bot.get_file(message.photo[-1].file_id)
                downloaded_file = bot.download_file(file_info.file_path)

                # –§–æ—Ä–º–∏—Ä—É–µ–º –ø—É—Ç—å –¥–ª—è —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è
                file_name = f"plan_{message.chat.id}_{message.message_id}.jpg"
                file_path = os.path.join(settings.MEDIA_ROOT, 'temp', file_name)

                # –°–æ–∑–¥–∞–µ–º –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏–∏ –µ—Å–ª–∏ –∏—Ö –Ω–µ—Ç
                os.makedirs(os.path.dirname(file_path), exist_ok=True)

                # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ñ–∞–π–ª
                with open(file_path, 'wb') as new_file:
                    new_file.write(downloaded_file)

                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –ø—É—Ç—å –≤ –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö
                pres.plan = file_path
                pres.save()
            else:
                bot.reply_to(message, "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –æ—Ç–ø—Ä–∞–≤—å—Ç–µ —Ñ–æ—Ç–æ–≥—Ä–∞—Ñ–∏—é –ø–ª–∞–Ω–∞")
                return get_plan(message.chat.id)
        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ –ø–ª–∞–Ω–∞: {e}")
            bot.reply_to(message, "–ü—Ä–æ–∏–∑–æ—à–ª–∞ –æ—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ –ø–ª–∞–Ω–∞. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ –µ—â–µ —Ä–∞–∑.")
            return get_plan(message.chat.id)
        return get_photo_outside(message.chat.id)

    def get_plan(id_):
        msg = bot.send_message(
            text="–û—Ç–ø—Ä–∞–≤—å—Ç–µ –ø–ª–∞–Ω –ø–æ–º–µ—â–µ–Ω–∏—è (–æ—Ç–ø—Ä–∞–≤—å—Ç–µ –∫–∞–∫ —Ñ–æ—Ç–æ)", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_plan)

    def register_type_rent(message):
        if check_exit(message):
            return
        pres.type_rent = message.text
        pres.save()
        return get_plan(message.chat.id)

    def get_type_rent(id_):
        msg = bot.send_message(
            text="–í—ã–±–µ—Ä–∏—Ç–µ —Ç–∏–ø –∞—Ä–µ–Ω–¥—ã –ø–æ–º–µ—â–µ–Ω–∏—è", chat_id=id_, reply_markup=TYPE_RENT
        )
        bot.register_next_step_handler(msg, register_type_rent)

    def register_rate(message):
        if check_exit(message):
            return
        pres.rate = message.text
        pres.save()
        return get_type_rent(message.chat.id)

    def get_rate(id_):
        msg = bot.send_message(
            text="–ñ–µ–ª–∞–µ–º–∞—è –∞—Ä–µ–Ω–¥–Ω–∞—è –ø–ª–∞—Ç–∞ –≤ –º–µ—Å—è—Ü, —Ä—É–±–ª–∏/–≤ –º–µ—Å—è—Ü (—É–∫–∞–∑—ã–≤–∞—Ç—å —Ç–æ–ª—å–∫–æ —Ü–∏—Ñ—Ä—ã)", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_rate)

    def register_height(message):
        if check_exit(message):
            return
        pres.height = message.text
        pres.save()
        return get_rate(message.chat.id)

    def get_height(id_):
        msg = bot.send_message(
            text="–í—ã—Å–æ—Ç–∞ –ø–æ—Ç–æ–ª–∫–æ–≤ –ø–æ–º–µ—â–µ–Ω–∏—è, –º–µ—Ç—Ä—ã (—É–∫–∞–∑—ã–≤–∞—Ç—å —Ç–æ–ª—å–∫–æ —Ü–∏—Ñ—Ä—É)", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_height)

    def register_floor(message):
        if check_exit(message):
            return
        pres.floor = message.text
        pres.save()
        return get_height(message.chat.id)

    def get_floor(id_):
        msg = bot.send_message(
            text="–í—ã–±–µ—Ä–∏—Ç–µ –Ω–∞ –∫–∞–∫–æ–º —ç—Ç–∞–∂–µ –≤–∞—à–µ –ø–æ–º–µ—â–µ–Ω–∏–µ", chat_id=id_, reply_markup=FLOOR_BUTTONS
        )
        bot.register_next_step_handler(msg, register_floor)

    def register_water_supply(message):
        if check_exit(message):
            return
        pres.water_supply = message.text
        pres.save()
        return get_floor(message.chat.id)

    def get_water_supply(id_):
        msg = bot.send_message(
            text="–ù–∞–ª–∏—á–∏–µ –≤–æ–¥–æ—Å–Ω–∞–±–∂–µ–Ω–∏—è –∏ –≤–æ–¥–æ–æ—Ç–≤–µ–¥–µ–Ω–∏—è –ø–æ–º–µ—â–µ–Ω–∏—è (–î–∞/–ù–µ—Ç)", chat_id=id_, reply_markup=WATER_SUPPLY
        )
        bot.register_next_step_handler(msg, register_water_supply)

    def register_power(message):
        if check_exit(message):
            return
        pres.power = message.text
        pres.save()
        return get_water_supply(message.chat.id)

    def get_power(id_):
        msg = bot.send_message(
            text="–ú–æ—â–Ω–æ—Å—Ç—å, –∫–í—Ç (—É–∫–∞–∑—ã–≤–∞—Ç—å —Ç–æ–ª—å–∫–æ —Ü–∏—Ñ—Ä—É)", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_power)

    def register_square(message):
        if check_exit(message):
            return
        pres.square = message.text
        pres.save()
        return get_power(message.chat.id)

    def get_square(id_):
        msg = bot.send_message(
            text="–ü–ª–æ—â–∞–¥—å –ø–æ–º–µ—â–µ–Ω–∏—è, –º–µ—Ç—Ä—ã –∫–≤–∞–¥—Ä–∞—Ç–Ω—ã–µ (—É–∫–∞–∑—ã–≤–∞—Ç—å —Ç–æ–ª—å–∫–æ —Ü–∏—Ñ—Ä—É)", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_square)

    def register_adress(message):
        if check_exit(message):
            return
        pres.adress = message.text
        pres.save()
        return get_square(message.chat.id)

    def get_adress(id_):
        msg = bot.send_message(
            text="–ê–¥—Ä–µ—Å –ø–æ–º–µ—â–µ–Ω–∏—è", chat_id=id_
        )
        bot.register_next_step_handler(msg, register_adress)

    def register_type_building(message):
        if check_exit(message):
            return

        pres.type_building = message.text
        pres.save()
        return get_adress(message.chat.id)

    def get_type_building(id_):
        msg = bot.send_message(
            text="–í—ã–±–µ—Ä–∏—Ç–µ —Ç–∏–ø –ø–æ–º–µ—â–µ–Ω–∏—è", chat_id=id_, reply_markup=TYPE_AREA
        )
        bot.register_next_step_handler(msg, register_type_building)

    get_type_building(call.message.chat.id)
    return


def show_all_presentations(call):
    pres = Presentations.objects.all()
    for pre in pres:
        doc = Document()

        doc.add_heading(f"{pre.square} –∫–≤.–º., {pre.adress}", level=1)
        doc.add_paragraph(f'–ó–¥–∞–Ω–∏–µ –ø–æ –∞–¥—Ä–µ—Å—É: {pre.adress}')
        doc.add_paragraph(f'–ü–ª–æ—â–∞–¥—å: {pre.square}')
        doc.add_paragraph(f'–≠–ª–µ–∫—Ç—Ä–æ—Å–Ω–∞–±–∂–µ–Ω–∏–µ –ø–æ–º–µ—â–µ–Ω–∏—è: {pre.power}')
        doc.add_paragraph(f'–í–æ–¥–æ—Å–Ω–∞–±–∂–µ–Ω–∏–µ –ø–æ–º–µ—â–µ–Ω–∏—è: {pre.water_supply}')
        doc.add_paragraph(f'–í—ã—Å–æ—Ç–∞ –ø–æ—Ç–æ–ª–∫–æ–≤: {pre.height}')
        doc.add_paragraph(f'–ñ–µ–ª–∞–µ–º–∞—è –∞—Ä–µ–Ω–¥–Ω–∞—è –ø–ª–∞—Ç–∞ –ø–æ–º–µ—â–µ–Ω–∏—è: {pre.rate}')
        doc.add_paragraph(f'–¢–∏–ø –∞—Ä–µ–Ω–¥—ã –ø–æ–º–µ—â–µ–Ω–∏—è: {pre.type_rent}')

        # –î–æ–±–∞–≤–ª—è–µ–º –ø–ª–∞–Ω –ø–æ–º–µ—â–µ–Ω–∏—è
        if pre.plan:
            try:
                plan_path = os.path.join(settings.MEDIA_ROOT, str(pre.plan))
                print(f"–ü—É—Ç—å –∫ –ø–ª–∞–Ω—É: {plan_path}")
                print(f"–°—É—â–µ—Å—Ç–≤—É–µ—Ç –ª–∏ —Ñ–∞–π–ª: {os.path.exists(plan_path)}")
                if os.path.exists(plan_path):
                    doc.add_heading('–ü–ª–∞–Ω –ø–æ–º–µ—â–µ–Ω–∏—è', level=2)
                    doc.add_picture(plan_path, width=Inches(6))
                else:
                    print(f"–§–∞–π–ª –ø–ª–∞–Ω–∞ –Ω–µ –Ω–∞–π–¥–µ–Ω –ø–æ –ø—É—Ç–∏: {plan_path}")
            except Exception as e:
                print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –¥–æ–±–∞–≤–ª–µ–Ω–∏–∏ –ø–ª–∞–Ω–∞: {e}")
                print(f"–ó–Ω–∞—á–µ–Ω–∏–µ pre.plan: {pre.plan}")
                print(f"–¢–∏–ø pre.plan: {type(pre.plan)}")
        else:
            print("–ü–ª–∞–Ω –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç –≤ –∑–∞–ø–∏—Å–∏")

        # –î–æ–±–∞–≤–ª—è–µ–º —Ñ–æ—Ç–æ —Å–Ω–∞—Ä—É–∂–∏
        if pre.photo_outside:
            try:
                outside_path = os.path.join(settings.MEDIA_ROOT, str(pre.photo_outside))
                if os.path.exists(outside_path):
                    doc.add_heading('–§–æ—Ç–æ —Å–Ω–∞—Ä—É–∂–∏', level=2)
                    doc.add_picture(outside_path, width=Inches(6))
            except Exception as e:
                print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –¥–æ–±–∞–≤–ª–µ–Ω–∏–∏ —Ñ–æ—Ç–æ —Å–Ω–∞—Ä—É–∂–∏: {e}")

        # –î–æ–±–∞–≤–ª—è–µ–º —Ñ–æ—Ç–æ –≤–Ω—É—Ç—Ä–∏
        if pre.photo_inside:
            try:
                inside_path = os.path.join(settings.MEDIA_ROOT, str(pre.photo_inside))
                if os.path.exists(inside_path):
                    doc.add_heading('–§–æ—Ç–æ –≤–Ω—É—Ç—Ä–∏', level=2)
                    doc.add_picture(inside_path, width=Inches(6))
            except Exception as e:
                print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –¥–æ–±–∞–≤–ª–µ–Ω–∏–∏ —Ñ–æ—Ç–æ –≤–Ω—É—Ç—Ä–∏: {e}")

        # –î–æ–±–∞–≤–ª—è–µ–º –ø—Ä–∏–º–µ—á–∞–Ω–∏—è
        if pre.additives and pre.additives != "0":
            doc.add_heading('–î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ –ø—Ä–∏–º–µ—á–∞–Ω–∏—è', level=2)
            doc.add_paragraph(pre.additives)

        # –°–æ—Ö—Ä–∞–Ω—è–µ–º –¥–æ–∫—É–º–µ–Ω—Ç —Å —É–Ω–∏–∫–∞–ª—å–Ω—ã–º –∏–º–µ–Ω–µ–º
        doc_name = f'presentation_{pre.user.user_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc_path = os.path.join(settings.MEDIA_ROOT, 'temp', doc_name)

        # –°–æ–∑–¥–∞–µ–º –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏—é –¥–ª—è –≤—Ä–µ–º–µ–Ω–Ω—ã—Ö —Ñ–∞–π–ª–æ–≤, –µ—Å–ª–∏ –µ—ë –Ω–µ—Ç
        os.makedirs(os.path.dirname(doc_path), exist_ok=True)

        doc.save(doc_path)

        # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –¥–æ–∫—É–º–µ–Ω—Ç –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é
        try:
            with open(doc_path, 'rb') as doc_file:
                pre.presentation = doc_path
                pre.save()
                bot.send_document(call.message.chat.id, doc_file, caption="–í–æ—Ç –≤–∞—à–∞ –ø—Ä–µ–∑–µ–Ω—Ç–∞—Ü–∏—è!")

        except Exception as e:
            print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –æ—Ç–ø—Ä–∞–≤–∫–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞: {e}")
        finally:
            # –£–¥–∞–ª—è–µ–º —Ñ–∞–π–ª –ø–æ—Å–ª–µ –æ—Ç–ø—Ä–∞–≤–∫–∏
            try:
                os.remove(doc_path)
            except Exception as e:
                print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —É–¥–∞–ª–µ–Ω–∏–∏ —Ñ–∞–π–ª–∞: {e}")


@bot.message_handler(commands=['presentations'])
def handle_presentations(message):
    """–û–±—Ä–∞–±–æ—Ç—á–∏–∫ –∫–æ–º–∞–Ω–¥—ã /presentations"""
    try:
        user = User.objects.filter(user_id=message.chat.id).first()
    except:
        print("–ü–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å –Ω–µ–∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω")
    try:
        pres = Presentations.objects.filter(user=user)
        if not pres.exists():
            bot.reply_to(message, "–ü–æ–∫–∞ –Ω–µ—Ç –Ω–∏ –æ–¥–Ω–æ–π –ø—Ä–µ–∑–µ–Ω—Ç–∞—Ü–∏–∏.")
            return

        for pre in pres:
            doc = Document()

            doc.add_heading(f'{pre.square} –∫–≤.–º., {pre.adress}', level=1)
            doc.add_paragraph(f'–°—Ç–∞–≤–∫–∞ –∞—Ä–µ–Ω–¥—ã: {pre.rate} —Ä—É–±.')
            doc.add_paragraph(f'–ü–ª–æ—â–∞–¥—å: {pre.square} –∫–≤.–º.')
            if pre.photo_outside:
                try:
                    outside_path = os.path.join(settings.MEDIA_ROOT, str(pre.photo_outside))
                    if os.path.exists(outside_path):
                        doc.add_heading('–§–æ—Ç–æ —Å–Ω–∞—Ä—É–∂–∏', level=2)
                        doc.add_picture(outside_path, width=Inches(6))
                except Exception as e:
                    print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –¥–æ–±–∞–≤–ª–µ–Ω–∏–∏ —Ñ–æ—Ç–æ —Å–Ω–∞—Ä—É–∂–∏: {e}")
            doc.add_heading("–•–∞—Ä–∞–∫—Ç–µ—Ä–∏—Å—Ç–∏–∫–∏")
            doc.add_paragraph(f'–¢–∏–ø –ø–æ–º–µ—â–µ–Ω–∏—è: {pre.type_building}')
            doc.add_paragraph(f'–†–∞—Å–ø–æ–ª–æ–∂–µ–Ω–∏–µ –ø–æ –∞–¥—Ä–µ—Å—É: {pre.adress}')
            doc.add_paragraph(f'–ü–ª–æ—â–∞–¥—å –ø–æ–º–µ—â–µ–Ω–∏—è: {pre.square} –∫–≤.–º.')
            doc.add_paragraph(f'–ú–æ—â–Ω–æ—Å—Ç—å: {pre.power} –∫–í—Ç')
            doc.add_paragraph(f'–í–æ–¥–æ—Å–Ω–∞–±–∂–µ–Ω–∏–µ –ø–æ–º–µ—â–µ–Ω–∏—è: {pre.water_supply}')
            doc.add_paragraph(f'–í—ã—Å–æ—Ç–∞ –ø–æ—Ç–æ–ª–∫–æ–≤: {pre.height} –º–µ—Ç—Ä–∞')
            doc.add_paragraph(f'–¢–∏–ø –∞—Ä–µ–Ω–¥—ã –ø–æ–º–µ—â–µ–Ω–∏—è: {pre.type_rent}')
            # –î–æ–±–∞–≤–ª—è–µ–º –ø–ª–∞–Ω –ø–æ–º–µ—â–µ–Ω–∏—è
            if pre.plan:
                try:
                    plan_path = os.path.join(settings.MEDIA_ROOT, str(pre.plan))
                    print(f"–ü—É—Ç—å –∫ –ø–ª–∞–Ω—É: {plan_path}")
                    print(f"–°—É—â–µ—Å—Ç–≤—É–µ—Ç –ª–∏ —Ñ–∞–π–ª: {os.path.exists(plan_path)}")
                    if os.path.exists(plan_path):
                        doc.add_heading('–ü–ª–∞–Ω –ø–æ–º–µ—â–µ–Ω–∏—è', level=2)
                        doc.add_picture(plan_path, width=Inches(6))
                    else:
                        print(f"–§–∞–π–ª –ø–ª–∞–Ω–∞ –Ω–µ –Ω–∞–π–¥–µ–Ω –ø–æ –ø—É—Ç–∏: {plan_path}")
                except Exception as e:
                    print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –¥–æ–±–∞–≤–ª–µ–Ω–∏–∏ –ø–ª–∞–Ω–∞: {e}")
                    print(f"–ó–Ω–∞—á–µ–Ω–∏–µ pre.plan: {pre.plan}")
                    print(f"–¢–∏–ø pre.plan: {type(pre.plan)}")
            else:
                print("–ü–ª–∞–Ω –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç –≤ –∑–∞–ø–∏—Å–∏")

                # –î–æ–±–∞–≤–ª—è–µ–º —Ñ–æ—Ç–æ –≤–Ω—É—Ç—Ä–∏
            if pre.photo_inside:
                try:
                    inside_path = os.path.join(settings.MEDIA_ROOT, str(pre.photo_inside))
                    if os.path.exists(inside_path):
                        doc.add_heading('–§–æ—Ç–æ –≤–Ω—É—Ç—Ä–∏', level=2)
                        doc.add_picture(inside_path, width=Inches(6))
                except Exception as e:
                    print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –¥–æ–±–∞–≤–ª–µ–Ω–∏–∏ —Ñ–æ—Ç–æ –≤–Ω—É—Ç—Ä–∏: {e}")

            # –î–æ–±–∞–≤–ª—è–µ–º –ø—Ä–∏–º–µ—á–∞–Ω–∏—è
            if pre.additives and pre.additives != "0":
                doc.add_heading('–î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω—ã–µ –ø—Ä–∏–º–µ—á–∞–Ω–∏—è', level=2)
                doc.add_paragraph(pre.additives)

            # –°–æ—Ö—Ä–∞–Ω—è–µ–º –¥–æ–∫—É–º–µ–Ω—Ç —Å —É–Ω–∏–∫–∞–ª—å–Ω—ã–º –∏–º–µ–Ω–µ–º
            doc_name = f'presentation_{pre.user.user_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            doc_path = os.path.join(settings.MEDIA_ROOT, 'temp', doc_name)

            # –°–æ–∑–¥–∞–µ–º –¥–∏—Ä–µ–∫—Ç–æ—Ä–∏—é –¥–ª—è –≤—Ä–µ–º–µ–Ω–Ω—ã—Ö —Ñ–∞–π–ª–æ–≤, –µ—Å–ª–∏ –µ—ë –Ω–µ—Ç
            os.makedirs(os.path.dirname(doc_path), exist_ok=True)

            doc.save(doc_path)

            # –û—Ç–ø—Ä–∞–≤–ª—è–µ–º –¥–æ–∫—É–º–µ–Ω—Ç –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é
            try:
                with open(doc_path, 'rb') as doc_file:
                    bot.send_document(message.chat.id, doc_file,
                                      caption=f"–ü—Ä–µ–∑–µ–Ω—Ç–∞—Ü–∏—è –¥–ª—è –ø–æ–º–µ—â–µ–Ω–∏—è –ø–æ –∞–¥—Ä–µ—Å—É: {pre.adress}")
            except Exception as e:
                print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –æ—Ç–ø—Ä–∞–≤–∫–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞: {e}")
                bot.reply_to(message, "–ü—Ä–æ–∏–∑–æ—à–ª–∞ –æ—à–∏–±–∫–∞ –ø—Ä–∏ –æ—Ç–ø—Ä–∞–≤–∫–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞.")
            finally:
                # –£–¥–∞–ª—è–µ–º —Ñ–∞–π–ª –ø–æ—Å–ª–µ –æ—Ç–ø—Ä–∞–≤–∫–∏
                try:
                    os.remove(doc_path)
                except Exception as e:
                    print(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —É–¥–∞–ª–µ–Ω–∏–∏ —Ñ–∞–π–ª–∞: {e}")

    except Exception as e:
        print(f"–û–±—â–∞—è –æ—à–∏–±–∫–∞: {e}")
        bot.reply_to(message, f"–ü—Ä–æ–∏–∑–æ—à–ª–∞ –æ—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ–∑–¥–∞–Ω–∏–∏ –ø—Ä–µ–∑–µ–Ω—Ç–∞—Ü–∏–π: {str(e)}")
