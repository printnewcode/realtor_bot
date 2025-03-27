from functools import wraps
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton
from docx import Document
from docx.shared import Inches
import os
from django.conf import settings
from datetime import datetime

from bot import bot
from bot.models import User, Presentations


def admin_permission(func):
    """
    Checking user for admin permission to access the function.
    """

    @wraps(func)
    def wrapped(message) -> None:
        user_id = message.from_user.id
        user = User.objects.get(user_id=user_id)
        if not user.is_admin:
            bot.send_message(user_id, '⛔ У вас нет администраторского доступа')
            return
        return func(message)

    return wrapped


@admin_permission
def admin_panel(message):
    """Админ-панель"""
    admin = User.objects.get(user_id=message.from_user.id)
    presentations = Presentations.objects.all().order_by('pk')
    reply = InlineKeyboardMarkup()
    try:
        for pres in presentations:
            reply.add(InlineKeyboardButton(text=f"Презентация {pres.pk}", callback_data=f"pres_{pres.pk}"))
        try:
            bot.send_message(text="Вот все презентации!", chat_id=message.chat.id, reply_markup=reply)
        except:
            bot.send_message(text="Презентаций нет!", chat_id=message.chat.id)
    except:
        bot.send_message(
            text="Упс, презентаций нет!",
            chat_id=message.chat.id
        )


def get_pres(call):
    _, data = call.data.split("_")
    presentation = Presentations.objects.filter(pk=data).first()
    pre = presentation
    doc = Document()

    doc.add_heading(f'{pre.square} кв.м., {pre.adress}', level=1)
    doc.add_paragraph(f'Ставка аренды: {pre.rate} руб./месяц')
    doc.add_paragraph(f'Площадь: {pre.square} кв.м.')
    if pre.photo_outside:
        try:
            outside_path = os.path.join(settings.MEDIA_ROOT, str(pre.photo_outside))
            if os.path.exists(outside_path):
                doc.add_heading('Фото снаружи', level=2)
                doc.add_picture(outside_path, width=Inches(6))
        except Exception as e:
            print(f"Ошибка при добавлении фото снаружи: {e}")
    doc.add_heading("Характеристики")
    doc.add_paragraph(f'Тип помещения: {pre.type_building}')
    doc.add_paragraph(f'Расположение по адресу: {pre.adress}')
    doc.add_paragraph(f'Площадь помещения: {pre.square} кв.м.')
    doc.add_paragraph(f'Мощность: {pre.power} кВт')
    doc.add_paragraph(f'Водоснабжение помещения: {pre.water_supply}')
    doc.add_paragraph(f'Высота потолков: {pre.height} метра')
    doc.add_paragraph(f'Тип аренды помещения: {pre.type_rent}')
    # Добавляем план помещения
    if pre.plan:
        try:
            plan_path = os.path.join(settings.MEDIA_ROOT, str(pre.plan))
            print(f"Путь к плану: {plan_path}")
            print(f"Существует ли файл: {os.path.exists(plan_path)}")
            if os.path.exists(plan_path):
                doc.add_heading('План помещения', level=2)
                doc.add_picture(plan_path, width=Inches(6))
            else:
                print(f"Файл плана не найден по пути: {plan_path}")
        except Exception as e:
            print(f"Ошибка при добавлении плана: {e}")
            print(f"Значение pre.plan: {pre.plan}")
            print(f"Тип pre.plan: {type(pre.plan)}")
    else:
        print("План отсутствует в записи")

        # Добавляем фото внутри
    if pre.photo_inside:
        try:
            inside_path = os.path.join(settings.MEDIA_ROOT, str(pre.photo_inside))
            if os.path.exists(inside_path):
                doc.add_heading('Фото внутри', level=2)
                doc.add_picture(inside_path, width=Inches(6))
        except Exception as e:
            print(f"Ошибка при добавлении фото внутри: {e}")

    # Добавляем примечания
    if pre.additives and pre.additives != "0":
        doc.add_heading('Дополнительные примечания', level=2)
        doc.add_paragraph(pre.additives)

    # Сохраняем документ с уникальным именем
    doc_name = f'presentation_{pre.user.user_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc_path = os.path.join(settings.MEDIA_ROOT, 'temp', doc_name)

    # Создаем директорию для временных файлов, если её нет
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)

    doc.save(doc_path)

    # Отправляем документ пользователю
    try:
        with open(doc_path, 'rb') as doc_file:
            bot.send_document(call.message.chat.id, doc_file,
                              caption=f"Презентация для помещения по адресу: {pre.adress}")
    except Exception as e:
        print(f"Ошибка при отправке документа: {e}")
        bot.reply_to(call.message, "Произошла ошибка при отправке документа.")
    finally:
        # Удаляем файл после отправки
        try:
            os.remove(doc_path)
        except Exception as e:
            print(f"Ошибка при удалении файла: {e}")
    bot.send_message(
        text=f"Контактные данные: {presentation.contact}\nID презентации: {presentation.pk}",
        chat_id=call.message.chat.id
    )


def admin_notification(presentation, id_):
    admin_list = User.objects.filter(is_admin=True)
    pre = presentation
    doc = Document()

    doc.add_heading(f'{pre.square} кв.м., {pre.adress}', level=1)
    doc.add_paragraph(f'Ставка аренды: {pre.rate} руб./месяц')
    doc.add_paragraph(f'Площадь: {pre.square} кв.м.')
    if pre.photo_outside:
        try:
            outside_path = os.path.join(settings.MEDIA_ROOT, str(pre.photo_outside))
            if os.path.exists(outside_path):
                doc.add_heading('Фото снаружи', level=2)
                doc.add_picture(outside_path, width=Inches(6))
        except Exception as e:
            print(f"Ошибка при добавлении фото снаружи: {e}")
    doc.add_heading("Характеристики")
    doc.add_paragraph(f'Тип помещения: {pre.type_building}')
    doc.add_paragraph(f'Расположение по адресу: {pre.adress}')
    doc.add_paragraph(f'Площадь помещения: {pre.square} кв.м.')
    doc.add_paragraph(f'Мощность: {pre.power} кВт')
    doc.add_paragraph(f'Водоснабжение помещения: {pre.water_supply}')
    doc.add_paragraph(f'Высота потолков: {pre.height} метра')
    doc.add_paragraph(f'Тип аренды помещения: {pre.type_rent}')
    # Добавляем план помещения
    if pre.plan:
        try:
            plan_path = os.path.join(settings.MEDIA_ROOT, str(pre.plan))
            print(f"Путь к плану: {plan_path}")
            print(f"Существует ли файл: {os.path.exists(plan_path)}")
            if os.path.exists(plan_path):
                doc.add_heading('План помещения', level=2)
                doc.add_picture(plan_path, width=Inches(6))
            else:
                print(f"Файл плана не найден по пути: {plan_path}")
        except Exception as e:
            print(f"Ошибка при добавлении плана: {e}")
            print(f"Значение pre.plan: {pre.plan}")
            print(f"Тип pre.plan: {type(pre.plan)}")
    else:
        print("План отсутствует в записи")

        # Добавляем фото внутри
    if pre.photo_inside:
        try:
            inside_path = os.path.join(settings.MEDIA_ROOT, str(pre.photo_inside))
            if os.path.exists(inside_path):
                doc.add_heading('Фото внутри', level=2)
                doc.add_picture(inside_path, width=Inches(6))
        except Exception as e:
            print(f"Ошибка при добавлении фото внутри: {e}")

    # Добавляем примечания
    if pre.additives and pre.additives != "0":
        doc.add_heading('Дополнительные примечания', level=2)
        doc.add_paragraph(pre.additives)

    # Сохраняем документ с уникальным именем
    doc_name = f'presentation_{pre.user.user_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc_path = os.path.join(settings.MEDIA_ROOT, 'temp', doc_name)

    # Создаем директорию для временных файлов, если её нет
    os.makedirs(os.path.dirname(doc_path), exist_ok=True)

    doc.save(doc_path)
    for admin in admin_list:
        try:
            with open(doc_path, 'rb') as doc_file:
                bot.send_document(admin.user_id, doc_file,
                                  caption=f"Презентация для помещения по адресу: {pre.adress}")
        except Exception as e:
            print(f"Ошибка при отправке документа: {e}")
        finally:
            # Удаляем файл после отправки
            try:
                os.remove(doc_path)
            except Exception as e:
                print(f"Ошибка при удалении файла: {e}")
        bot.send_message(
            text=f"Пользователь @{presentation.user.username} создал презентацию.\n\nКонтактные данные: {presentation.contact}\n\nID: {presentation.pk}",
            chat_id=id_
        )
